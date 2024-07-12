import markdown
from docx import Document
from docx.shared import Inches
from bs4 import BeautifulSoup


# Function to add headings to the document
def add_heading(document, text, level):
    if level == 1:
        document.add_heading(text, level=1)
    elif level == 2:
        document.add_heading(text, level=2)
    # Add more conditions for other levels of headings if needed


# Function to convert Markdown to Word Document
def markdown_to_word(md_file, docx_file):
    with open(md_file, 'r') as f:
        text = f.read()
        html = markdown.markdown(text)

    document = Document()
    soup = BeautifulSoup(html, features='html.parser')

    # Parse the HTML and add elements to the Word document
    for tag in soup.find_all():
        if tag.name == 'h1':
            add_heading(document, tag.text, level=1)
        elif tag.name == 'h2':
            add_heading(document, tag.text, level=2)
        # Add more conditions for other tags (paragraphs, images etc.) if needed
    document.save(docx_file)


# Main function to run the conversion process
def main():
    md_file = "input.md"  # Path to your markdown file
    docx_file = "output.docx"  # Output Word Document file name
    markdown_to_word(md_file, docx_file)


if __name__ == '__main__':
    main()
